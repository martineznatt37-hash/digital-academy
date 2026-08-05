# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

outs = [
    r"C:\Users\marti\OneDrive\Desktop\digital-academy\docs\DOCUMENTACION-ESCAPE-DEL-HACKER.docx",
    r"C:\Users\marti\OneDrive\Desktop\Escape-Del-Hacker\docs\DOCUMENTACION-ESCAPE-DEL-HACKER.docx",
]


def set_run_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading_custom(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, "Calibri", 16 if level == 1 else 13, True, RGBColor(0, 102, 128))
    return p


def add_p(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(it, style="List Bullet")
        for run in p.runs:
            set_run_font(run, size=11)


def add_numbered(doc, items):
    for it in items:
        p = doc.add_paragraph(it, style="List Number")
        for run in p.runs:
            set_run_font(run, size=11)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, bold=True, size=11, color=RGBColor(0, 80, 100))
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            table.rows[r + 1].cells[c].text = val
            for p in table.rows[r + 1].cells[c].paragraphs:
                for run in p.runs:
                    set_run_font(run, size=10)
    doc.add_paragraph()


doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("UNIVERSIDAD AZTECA")
set_run_font(r, size=14, bold=True, color=RGBColor(0, 90, 110))

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run("Ingeniería en Sistemas Computacionales")
set_run_font(r2, size=12, color=RGBColor(60, 60, 60))

t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run("\nDOCUMENTACIÓN TÉCNICA DEL PROYECTO")
set_run_font(r3, size=18, bold=True, color=RGBColor(0, 120, 140))

t4 = doc.add_paragraph()
t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = t4.add_run("ESCAPE DEL HACKER 2.0")
set_run_font(r4, size=22, bold=True, color=RGBColor(0, 160, 180))

t5 = doc.add_paragraph()
t5.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = t5.add_run("Interfaz gráfica / Juego web educativo de ciberseguridad")
set_run_font(r5, size=12)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
rm = meta.add_run("\nTipo: Interfaz Gráfica\nPeriodo: 26-3")
set_run_font(rm, size=11, color=RGBColor(80, 80, 80))

doc.add_page_break()

add_heading_custom(doc, "1. ¿Qué es el proyecto?", 1)
add_p(
    doc,
    "Escape del Hacker es un juego web donde el jugador es teletransportado a un sistema hackeado y debe escapar nivel por nivel. En cada nivel enfrenta un tipo de malware, recoge llaves, evita enemigos y llega a la puerta de salida. Algunos niveles incluyen pelea contra un jefe.",
)
add_p(doc, "El objetivo académico es demostrar:", bold=True)
add_bullets(
    doc,
    [
        "Interfaz atractiva y usable en PC y móvil",
        "Funcionalidades completas (registro, juego, progreso)",
        "Arquitectura escalable (frontend + API + datos)",
        "Código documentado y demostración clara",
    ],
)

add_heading_custom(doc, "2. Problema / necesidad", 1)
add_p(
    doc,
    "Muchas personas conocen palabras como virus o ransomware, pero no visualizan cómo “se siente” estar dentro de un sistema atacado. Este juego:",
)
add_numbered(
    doc,
    [
        "Enseña conceptos de malware de forma lúdica (un malware por nivel).",
        "Motiva con personajes, habilidades y puntuación.",
        "Funciona en cualquier dispositivo con un enlace (Netlify).",
        "Guarda cuentas y progreso con un servidor en la nube (Render).",
    ],
)

add_heading_custom(doc, "3. Qué se construyó (resumen de trabajo)", 1)
add_table(
    doc,
    ["Etapa", "Qué se hizo"],
    [
        ["Concepto", "Tema ciberseguridad / infiltración"],
        ["Interfaz", "Login, menú, selección de personajes, HUD, intros"],
        ["Gameplay", "10 niveles, llaves, muros, enemigos, jefes, habilidades"],
        ["Arte", "Sprites de personajes, malware, NPC “El Hacker”"],
        ["Audio", "Música/SFX tipo chiptune"],
        ["Responsive", "Controles táctiles en teléfono; teclado en PC"],
        ["Backend", "API registro/login/guardar progreso"],
        ["Deploy", "Frontend en Netlify, API en Render, código en GitHub"],
    ],
)

add_heading_custom(doc, "4. Arquitectura del sistema", 1)
add_p(doc, "Flujo del sistema:", bold=True)
add_p(
    doc,
    "Jugador (PC / Celular / Tablet) → Netlify (frontend: escape-public) → Render (backend Node/Express /api/escape) → SQLite (tabla escape_players)",
)
add_heading_custom(doc, "4.1 Separación de responsabilidades", 2)
add_bullets(
    doc,
    [
        "Frontend: dibuja el juego, lee teclado/touch, llama a la API.",
        "Backend: valida usuarios, hashea contraseñas, guarda progreso.",
        "Base de datos: tabla escape_players (usuario, hash, récord, nivel máximo).",
    ],
)
add_p(doc, "Esto permite crecer: nuevos niveles en el cliente, nuevas rutas en el servidor, sin rehacer todo.")

add_heading_custom(doc, "5. Estructura de archivos del juego", 1)
add_p(doc, "Escape-Del-Hacker/")
add_bullets(
    doc,
    [
        "juego/ → código fuente (HTML + assets)",
        "escape-public/ → listo para Netlify",
        "server/ → API para Render (routes/escape.js = login / registro / save)",
        "docs/ → documentación y presentación",
    ],
)
add_p(doc, "Archivos clave:", bold=True)
add_bullets(
    doc,
    [
        "escape-del-hacker.html — UI + lógica + canvas",
        "server/routes/escape.js — API de cuentas",
        "server/server.js — arranque Express",
        "assets/ — gráficos",
    ],
)

add_heading_custom(doc, "6. Funcionalidades principales", 1)
add_numbered(
    doc,
    [
        "Registro / Login con usuario y contraseña",
        "Selección de personaje (habilidad y arma distintas)",
        "10 niveles con malware y tema visual propio",
        "Mecánicas: llaves, puerta, láseres, firewalls, código fantasma, jefes",
        "Controles: teclado (PC) o D-pad + habilidad (móvil)",
        "Progreso: récord y nivel máximo sincronizable por API",
        "Publicación web accesible desde cualquier dispositivo",
    ],
)
add_p(doc, "En la rúbrica, estos puntos equivalen a las “funciones clave” del sistema.")

add_heading_custom(doc, "7. Escalabilidad", 1)
add_p(doc, "Cómo se puede ampliar sin reescribir el núcleo:")
add_table(
    doc,
    ["Idea nueva", "Dónde se agrega"],
    [
        ["Nivel 11+", "Objeto en arreglo niveles[]"],
        ["Personaje nuevo", "Objeto en personajes[] + PNG"],
        ["Ranking global", "Nueva ruta /api/escape/ranking + tabla"],
        ["Tienda de skins", "Nuevo módulo API + UI"],
        ["App Android", "Capacitor sobre el mismo frontend"],
    ],
)

add_heading_custom(doc, "8. Tecnologías usadas", 1)
add_bullets(
    doc,
    [
        "Frontend: HTML5, CSS3, JavaScript (Canvas 2D)",
        "Backend: Node.js, Express, SQLite (node:sqlite)",
        "Seguridad: bcrypt, JWT",
        "Deploy: Netlify (estático), Render (API), GitHub (versión)",
    ],
)

add_heading_custom(doc, "9. Cómo publicar / demostrar", 1)
add_numbered(
    doc,
    [
        "Subir carpeta escape-public a Netlify.",
        "Tener el servidor en Render.",
        "Abrir: https://TU-SITIO.netlify.app/?api=https://escape-del-hacker.onrender.com/api/escape",
        "Registrarse → jugar un nivel → entrar desde otro dispositivo.",
    ],
)
add_p(doc, "Detalle paso a paso también en el archivo COMO-PUBLICAR-ESCAPE.md.")

add_heading_custom(doc, "10. Beneficios para el usuario final", 1)
add_bullets(
    doc,
    [
        "Jugar sin instalar programas pesados (solo navegador).",
        "Continuar partida/cuenta en otro aparato.",
        "Aprender tipos de malware de forma visual.",
        "Experiencia con controles claros en PC y celular.",
    ],
)

add_heading_custom(doc, "11. Conclusión", 1)
add_p(
    doc,
    "Se entregó un sistema completo de interfaz gráfica con juego funcional, despliegue multiplataforma y backend de cuentas. La arquitectura cliente–servidor y los datos configurables (niveles/personajes) respaldan la escalabilidad exigida en la rúbrica.",
)

firma = doc.add_paragraph()
runf = firma.add_run("\n\n______________________________\nNombre del estudiante / Firma")
set_run_font(runf, size=11)

for path in outs:
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)
    print("OK", path)
